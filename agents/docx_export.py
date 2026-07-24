"""
Writer 에이전트 DOCX 내보내기

체크포인트3 요구사항("제공된 기획방법론 템플릿(KOSENA 등) 준수, DOCX/PPTX 내보내기 지원")과
KOSENA 강의자료(제출 형식: 본문 PDF/DOCX, A4)를 근거로 추가한 모듈.

writer.py의 run_writer()가 마크다운 문자열을 조립하는 것과 별개로, 이 모듈은 같은
구조화된 데이터(MarketSizing, pestel_summaries, Competitor 리스트, 텍스트 섹션들)를
그대로 받아 Word 문서 요소(제목, 표, 목록)로 직접 변환한다.

2026-07-23: 사용자가 업로드한 참고 docx(`웨어러블_헬스케어_기기_북미_시니어_건강관리_기획서.docx`)의
실제 OOXML(word/document.xml)을 직접 열어 색상·크기·테두리 값을 1:1로 대조해 재작성했다
(파이프라인 문서 36절). 이전에 참고했던 `docx_design.md` 프로즈 설명은 실제 파일과 여러
군데 어긋나 있었음이 이번 대조로 확인됐다 — 표 테두리(전체그리드 아님, 상하만), 각주
이탤릭 여부, 푸터 구성, H1/H2 색 분리, PESTEL 배지 범위, 페이지 나눔 유무 등. 이번 버전은
프로즈 설명이 아니라 실제 파일의 XML 값을 근거로 삼았다.
"""

import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from fact_store.schema import Competitor, CompetitorType, MarketSizing
from fact_store.store import get_fact

# ---------------------------------------------------------------------------
# 1. 컬러 팔레트 · 폰트 (참고 docx의 word/document.xml 실측값 기준)
# ---------------------------------------------------------------------------
FONT_NAME = "맑은 고딕"

PRIMARY_HEX = "0F4C5C"
ACCENT_HEX = "1B8A8F"
SUB_HEX = "5A6B73"
BODY_HEX = "222222"
WARN_HEX = "9A6B00"
WARN_BG_HEX = "FBF6E9"
ACCENT2_HEX = "E9F4F4"
LIGHT_HEX = "F5F8F8"
LINE_HEX = "C9DADB"
HEAD_BG_HEX = PRIMARY_HEX
FOOTER_GRAY_HEX = "8A9AA0"

COLOR_PRIMARY = RGBColor(0x0F, 0x4C, 0x5C)
COLOR_ACCENT = RGBColor(0x1B, 0x8A, 0x8F)
COLOR_SUB = RGBColor(0x5A, 0x6B, 0x73)
COLOR_BODY = RGBColor(0x22, 0x22, 0x22)
COLOR_WARN = RGBColor(0x9A, 0x6B, 0x00)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_FOOTER_GRAY = RGBColor(0x8A, 0x9A, 0xA0)

_AXIS_KOREAN = {
    "Political": "정치적 환경",
    "Economic": "경제적 환경",
    "Social": "사회적 환경",
    "Technological": "기술적 환경",
    "Environmental": "환경적 요인",
    "Legal": "법·규제 환경",
}

_CM_PER_TWIP = 1 / 566.929

# 경쟁사 표: 경쟁사/가격/핵심기능/타깃·비고 4컬럼 (twips: 1650/1700/4456/1500)
_COMPETITOR_TABLE_HEADERS = ["경쟁사", "가격", "핵심 기능 (차별화 포인트·요약)", "타깃·비고"]
_COMPETITOR_TABLE_COL_WIDTHS_CM = [w * _CM_PER_TWIP for w in (1650, 1700, 4456, 1500)]

_COMPETITOR_TYPE_ORDER = [CompetitorType.DIRECT, CompetitorType.INDIRECT, CompetitorType.POTENTIAL]

# 시장 규모 데이터 표: 구분/규모(USD)/산정 방식/정의 (twips: 1500/2300/1600/4106)
_SIZING_TABLE_COL_WIDTHS_CM = [w * _CM_PER_TWIP for w in (1500, 2300, 1600, 4106)]
_SIZING_DEFINITIONS = {
    "TAM": "시장 전체 규모",
    "SAM": "TAM 중 유효 시장 세그먼트",
    "SOM": "신규 진입자의 초기 확보 목표 시장",
}

# 출처 부록 표: 등급/판정 / 수집 fact / 출처 (twips: 1150/5456/2700)
_SOURCE_TABLE_COL_WIDTHS_CM = [w * _CM_PER_TWIP for w in (1150, 5456, 2700)]

# 숫자·비율·개수 등 핵심 수치를 본문 안에서 부분 강조할 때 쓰는 패턴
_EMPHASIS_RE = re.compile(
    r"\$?[\d,]+(?:\.\d+)?(?:천만|억|만)?%?|\d+(?:\.\d+)?%|\d+(?:개|가지|건)(?:사)?"
)


def _fmt_num(value: Optional[float], unit: str) -> str:
    if value is None:
        return "계산 불가"
    return f"{value:,.0f} {unit}"


_UNIT_KOREAN = {"USD": "달러", "KRW": "원", "EUR": "유로", "JPY": "엔"}


def _fmt_compact(value: Optional[float], unit: str = "", dollar_prefix: bool = False) -> str:
    """0이 많은 원 숫자를 억/만 단위로 압축 표기한다 (참고 docx 실측: '$153.6억',
    '9,216만', '153.6억 달러' 등 — 1억(1e8) 이상은 억 단위 소수1자리, 1만(1e4)
    이상은 만 단위 정수로 축약). dollar_prefix=True면 '$153.6억'처럼 $ 접두사만
    붙이고(표지 KPI 카드), False면 단위 한글명을 뒤에 붙인다(데이터 표)."""
    if value is None:
        return "계산 불가"
    sign = "-" if value < 0 else ""
    v = abs(value)
    if v >= 1e8:
        body = f"{v / 1e8:,.1f}억"
    elif v >= 1e4:
        body = f"{v / 1e4:,.0f}만"
    else:
        body = f"{v:,.0f}"
    if dollar_prefix:
        return f"{sign}${body}"
    word = _UNIT_KOREAN.get(unit, unit or "")
    return f"{sign}{body}" + (f" {word}" if word else "")


# ---------------------------------------------------------------------------
# 저수준 헬퍼: 폰트, 셀 배경/테두리
# ---------------------------------------------------------------------------
def _set_run_font(
    run,
    size: Optional[Pt] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[RGBColor] = None,
) -> None:
    """맑은 고딕을 ascii/eastAsia/hAnsi 모두에 지정 — python-docx의 run.font.name만
    설정하면 한글은 Word 기본 동아시아 폰트로 렌더링돼 실제로는 안 바뀐다."""
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_run(paragraph, text: str, **kwargs):
    run = paragraph.add_run(text)
    _set_run_font(run, **kwargs)
    return run


def _shade_run(run, hex_color: str) -> None:
    """런(문단 내 텍스트 조각) 배경색 — PESTEL 배지("P" 한 글자)에 사용."""
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_hv_borders(cell, top_color: str = LINE_HEX, top_size: int = 4, bottom_color: str = LINE_HEX, bottom_size: int = 4) -> None:
    """참고 docx 실측 결과 — 표는 좌우 테두리가 없고 상/하 얇은 선만 있다(잡지형 미니멀
    스타일). KPI 카드처럼 위쪽만 강조색 굵은 선을 쓰는 경우는 top_color/top_size로 조정."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(top_size))
    top.set(qn("w:color"), top_color)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "nil")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(bottom_size))
    bottom.set(qn("w:color"), bottom_color)
    right = OxmlElement("w:right")
    right.set(qn("w:val"), "nil")
    for el in (top, left, bottom, right):
        borders.append(el)
    tcPr.append(borders)


def _set_cell_margins(cell, top_twip: int = 70, bottom_twip: int = 70, left_twip: int = 110, right_twip: int = 110) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top_twip), ("bottom", bottom_twip), ("start", left_twip), ("end", right_twip)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _apply_table_grid_fallback(table) -> None:
    """표 전체의 기본(fallback) 테두리 — 개별 셀에서 top/bottom(및 none 처리한
    left/right)을 항상 지정하므로 실제로는 이 표준 그리드가 눈에 보이지 않지만,
    참고 docx도 이 fallback을 tblPr에 남겨두고 있어 동일하게 둔다."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _set_table_grid_widths(table, col_widths_cm: list[float]) -> None:
    """tblGrid/tblW를 명시적으로 지정한다. python-docx는 cell.width만 설정해도
    autofit=False(=tblLayout fixed)로는 되지만, 초기 tblGrid의 컬럼 폭까지 갱신하지
    않아 LibreOffice/Word가 좁은 컬럼(예: 종합 시사점 번호 칸)을 실제보다 훨씬 넓게
    렌더링하는 문제가 있다 — 표 생성 직후 이 함수로 grid 폭을 다시 씌워준다."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    total_twips = int(sum(col_widths_cm) * 566.929)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(total_twips))
    tblPr.append(tblW)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width_cm in col_widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width_cm * 566.929)))
        grid.append(col)
    tbl.insert(list(tbl).index(tblPr) + 1, grid)


def _mark_header_row(row) -> None:
    """여러 페이지에 걸치는 표에서 헤더 행이 반복되도록 표시."""
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    trPr.append(el)


def _set_a4(doc: Document) -> None:
    """KOSENA 제출 형식(A4) 준수 + 참고 docx 실측 여백(상하 1200twip/좌우 1300twip)."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Pt(1200 / 20)
    section.bottom_margin = Pt(1200 / 20)
    section.left_margin = Pt(1300 / 20)
    section.right_margin = Pt(1300 / 20)


def _add_footer(doc: Document, footer_text: str) -> None:
    """푸터: 문서제목 — 페이지 번호 (참고 docx 실측: word/footer1.xml)."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, f"{footer_text}  |  ", size=Pt(7), color=COLOR_FOOTER_GRAY)

    run = p.add_run()
    _set_run_font(run, size=Pt(7), color=COLOR_FOOTER_GRAY)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)


def _split_emphasis_prefix(text: str) -> tuple[str, str]:
    """'핵심값 — 설명' 또는 '[태그] KEY=VALUE (설명)' 형태에서 앞쪽 핵심 부분과
    뒤쪽 설명 부분을 분리한다 (참고 docx 실측: 불릿의 앞부분만 bold+primary)."""
    if " — " in text:
        key, _, rest = text.partition(" — ")
        return key + "  —  ", rest
    if "(" in text:
        idx = text.index("(")
        return text[:idx], text[idx:]
    return text, ""


def _add_bullets(doc: Document, items: list[str]) -> None:
    """가정·근거 등 핵심 불릿: 앞부분(핵심값)만 bold+primary, 뒷부분은 일반 텍스트."""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        head, rest = _split_emphasis_prefix(item)
        _add_run(p, head, bold=True, size=Pt(10), color=COLOR_PRIMARY)
        if rest:
            _add_run(p, rest, size=Pt(10), color=COLOR_BODY)


def _add_h1(doc: Document, number: int, title: str) -> None:
    """H1: '01  제목' — 번호는 accent, 제목 텍스트는 primary, 하단 accent 밑줄
    (참고 docx 실측: w:pBdr bottom color=1B8A8F sz=12 space=4)."""
    p = doc.add_paragraph(style="Heading 1")
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:color"), ACCENT_HEX)
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    _add_run(p, f"{number:02d}  ", size=Pt(15), bold=True, color=COLOR_ACCENT)
    _add_run(p, title, size=Pt(15), bold=True, color=COLOR_PRIMARY)


def _add_h2(doc: Document, text: str) -> None:
    """H2: '▎ 소제목' — 마크는 accent, 텍스트는 primary."""
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    _add_run(p, "▎", size=Pt(12), bold=True, color=COLOR_ACCENT)
    _add_run(p, f" {text}", size=Pt(12), bold=True, color=COLOR_PRIMARY)


def _add_body_paragraph(doc: Document, text: str, emphasize: bool = False):
    """일반 본문 10pt(222222). emphasize=True면 숫자·비율·개수 등 핵심 수치만
    부분적으로 bold+primary 강조 (참고 docx 실측: 개요 문단이 이 패턴)."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(15)
    if not emphasize:
        _add_run(p, text, size=Pt(10), color=COLOR_BODY)
        return p
    pos = 0
    for m in _EMPHASIS_RE.finditer(text):
        if m.start() > pos:
            _add_run(p, text[pos:m.start()], size=Pt(10), color=COLOR_BODY)
        _add_run(p, m.group(), bold=True, size=Pt(10), color=COLOR_PRIMARY)
        pos = m.end()
    if pos < len(text):
        _add_run(p, text[pos:], size=Pt(10), color=COLOR_BODY)
    return p


def _add_note(doc: Document, text: str) -> None:
    """각주성 문구(※...): 8pt, sub color, 이탤릭 (참고 docx 실측: w:i 있음)."""
    p = doc.add_paragraph()
    _add_run(p, text, size=Pt(8), italic=True, color=COLOR_SUB)


def _add_warn_box(doc: Document, text: str) -> None:
    """콜아웃(⚠): FBF6E9 배경 + 9A6B00 Bold."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    _apply_table_grid_fallback(table)
    cell = table.rows[0].cells[0]
    cell.width = Cm(16.4)
    _shade_cell(cell, WARN_BG_HEX)
    _set_cell_hv_borders(cell)
    _set_cell_margins(cell)
    p = cell.paragraphs[0]
    _add_run(p, f"⚠ {text}", bold=True, size=Pt(9.5), color=COLOR_WARN)


# ---------------------------------------------------------------------------
# 2. 표지 — 뒤이어 페이지 나눔 없이 01 개요로 바로 이어진다 (참고 docx 실측)
# ---------------------------------------------------------------------------
def _add_cover(doc: Document, topic: str, target_market: str, created_date: str) -> None:
    label = doc.add_paragraph()
    _add_run(label, "사업 기획서 초안  |  DRAFT", size=Pt(9), bold=True, color=COLOR_ACCENT)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(6)
    _add_run(title_p, topic, size=Pt(26), bold=True, color=COLOR_PRIMARY)

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(18)
    _add_run(subtitle_p, f"{target_market} 시장 진입 기획", size=Pt(15), color=COLOR_SUB)

    info_rows = [
        ("목표 시장", target_market),
        ("작성일", created_date),
        ("작성", "AI 서비스 기획 보조 Multi-Agent — 자동 생성 초안, 사람 검수 필요"),
    ]
    # 2200:7106 twip = 라벨 24% : 값 76% (A4 기준). 격자 스타일을 절대 쓰지 않고
    # (table.style=None) 좌우 테두리는 nil로 지워, 위/아래 얇은 회색선만 남긴다.
    # tblGrid/tblW를 명시적으로 지정해야 렌더러가 좁은 라벨 열(24%)을 실제로 지킨다.
    label_width_cm = 2200 * _CM_PER_TWIP
    value_width_cm = 7106 * _CM_PER_TWIP
    table = doc.add_table(rows=len(info_rows), cols=2)
    table.style = None
    table.autofit = False
    _set_table_grid_widths(table, [label_width_cm, value_width_cm])
    for row, (label_text, value_text) in zip(table.rows, info_rows):
        label_cell, value_cell = row.cells
        label_cell.width = Cm(label_width_cm)
        value_cell.width = Cm(value_width_cm)
        _shade_cell(label_cell, ACCENT2_HEX)
        _set_cell_hv_borders(label_cell)
        _set_cell_hv_borders(value_cell)
        # 상하 0.05in(=72twip), 좌우 0.08in(=115twip)
        _set_cell_margins(label_cell, top_twip=72, bottom_twip=72, left_twip=115, right_twip=115)
        _set_cell_margins(value_cell, top_twip=72, bottom_twip=72, left_twip=115, right_twip=115)
        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        label_p = label_cell.paragraphs[0]
        label_p.paragraph_format.space_before = Pt(0)
        label_p.paragraph_format.space_after = Pt(0)
        _add_run(label_p, label_text, bold=True, size=Pt(9.5), color=COLOR_PRIMARY)

        value_p = value_cell.paragraphs[0]
        value_p.paragraph_format.space_before = Pt(0)
        value_p.paragraph_format.space_after = Pt(0)
        _add_run(value_p, value_text, size=Pt(9.5), color=COLOR_BODY)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# 3. 01 개요 — TAM·SAM·SOM 카드
# ---------------------------------------------------------------------------
def _add_kpi_cards(doc: Document, market_sizing: Optional[MarketSizing]) -> None:
    if market_sizing is None:
        _add_body_paragraph(doc, "시장조사 에이전트 실행 이력이 없어 시장 규모 데이터가 없습니다.")
        return

    cards = [
        ("TAM", _fmt_compact(market_sizing.tam_topdown, dollar_prefix=True), "2022년 시장 규모 기준 (Top-down)"),
        ("SAM", _fmt_compact(market_sizing.sam_topdown, dollar_prefix=True), "TAM 중 유효 시장 세그먼트"),
        ("SOM", _fmt_compact(market_sizing.som_topdown, dollar_prefix=True), "초기 확보 가능 목표 (SAM 대비 비중)"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    _apply_table_grid_fallback(table)
    row = table.rows[0]
    for cell, (label_text, value_text, caption_text) in zip(row.cells, cards):
        cell.width = Cm(3102 * _CM_PER_TWIP)
        _shade_cell(cell, LIGHT_HEX)
        _set_cell_hv_borders(cell, top_color=ACCENT_HEX, top_size=16, bottom_color=LINE_HEX, bottom_size=4)
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        label_p = cell.paragraphs[0]
        label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(label_p, label_text, bold=True, size=Pt(9), color=COLOR_ACCENT)

        value_p = cell.add_paragraph()
        value_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(value_p, value_text, bold=True, size=Pt(15), color=COLOR_PRIMARY)

        caption_p = cell.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(caption_p, caption_text, size=Pt(7.5), color=COLOR_SUB)


# ---------------------------------------------------------------------------
# 4. 02 시장 규모 분석 — 데이터 표 (구분/규모/산정방식/정의)
# ---------------------------------------------------------------------------
def _add_market_sizing_table(doc: Document, market_sizing: MarketSizing) -> None:
    unit = market_sizing.unit
    rows_data = [
        ("TAM", _fmt_compact(market_sizing.tam_topdown, unit), "Top-down", _SIZING_DEFINITIONS["TAM"]),
        ("SAM", _fmt_compact(market_sizing.sam_topdown, unit), "Top-down", _SIZING_DEFINITIONS["SAM"]),
        ("SOM", _fmt_compact(market_sizing.som_topdown, unit), "Top-down", _SIZING_DEFINITIONS["SOM"]),
    ]
    if market_sizing.som_bottomup:
        rows_data.append(
            ("SOM", _fmt_compact(market_sizing.som_bottomup, unit), "Bottom-up", _SIZING_DEFINITIONS["SOM"])
        )

    headers = ["구분", f"규모 ({market_sizing.unit})", "산정 방식", "정의"]
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    _apply_table_grid_fallback(table)
    header_row = table.rows[0]
    _mark_header_row(header_row)
    for cell, header_text, width in zip(header_row.cells, headers, _SIZING_TABLE_COL_WIDTHS_CM):
        cell.width = Cm(width)
        _shade_cell(cell, HEAD_BG_HEX)
        _set_cell_hv_borders(cell)
        _set_cell_margins(cell)
        _add_run(cell.paragraphs[0], header_text, bold=True, size=Pt(9.5), color=COLOR_WHITE)

    for label, value, method, definition in rows_data:
        row_cells = table.add_row().cells
        for cell, width in zip(row_cells, _SIZING_TABLE_COL_WIDTHS_CM):
            cell.width = Cm(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            _set_cell_hv_borders(cell)
            _set_cell_margins(cell)
        _add_run(row_cells[0].paragraphs[0], label, bold=True, size=Pt(9.5), color=COLOR_PRIMARY)
        _add_run(row_cells[1].paragraphs[0], value, bold=True, size=Pt(9.5), color=COLOR_PRIMARY)
        _add_run(row_cells[2].paragraphs[0], method, size=Pt(9.5), color=COLOR_BODY)
        _add_run(row_cells[3].paragraphs[0], definition, size=Pt(9.5), color=COLOR_BODY)


# ---------------------------------------------------------------------------
# 5. 03 PESTEL — 배지("P" 한 글자만 배경색, 나머지는 배경 없음)
# ---------------------------------------------------------------------------
def _add_pestel_axis(doc: Document, axis: str, narrative: str, excluded: list) -> None:
    letter = axis[0].upper() if axis else "?"
    korean = _AXIS_KOREAN.get(axis, "")

    heading_p = doc.add_paragraph(style="Heading 2")
    heading_p.paragraph_format.space_before = Pt(11)
    heading_p.paragraph_format.space_after = Pt(5)

    badge_run = heading_p.add_run(f" {letter} ")
    _set_run_font(badge_run, bold=True, size=Pt(11), color=COLOR_WHITE)
    _shade_run(badge_run, ACCENT_HEX)

    rest_text = f"  {axis} — {korean}" if korean else f"  {axis}"
    _add_run(heading_p, rest_text, bold=True, size=Pt(11), color=COLOR_PRIMARY)

    _add_body_paragraph(doc, narrative)
    if excluded:
        _add_note(doc, f"※ 중요도가 낮아 요약에서 제외한 fact: {len(excluded)}건")


# ---------------------------------------------------------------------------
# 6. 04 경쟁사 분석 — 표
# ---------------------------------------------------------------------------
def _feature_cell_text(cell, key_features: list[str]) -> None:
    """핵심 기능 셀: 1행 차별화 포인트(bold, primary, 9pt), 2행 나머지 기능 · 로 연결."""
    if not key_features:
        _add_run(cell.paragraphs[0], "(정보 없음)", size=Pt(9), color=COLOR_BODY)
        return
    headline, *rest = key_features
    p1 = cell.paragraphs[0]
    _add_run(p1, headline, bold=True, size=Pt(9), color=COLOR_PRIMARY)
    if rest:
        p2 = cell.add_paragraph()
        _add_run(p2, " · ".join(rest), size=Pt(8), color=COLOR_BODY)


def _target_remark_cell_text(cell, competitor: Competitor) -> None:
    """타깃·비고: target_customer 기본 + channel/funding_or_revenue 값 있는 경우만 추가."""
    parts = []
    if competitor.target_customer:
        parts.append(competitor.target_customer)
    extra = [v for v in (competitor.funding_or_revenue, competitor.channel) if v]
    parts.extend(extra)
    text = " / ".join(parts) if parts else "(정보 없음)"
    _add_run(cell.paragraphs[0], text, size=Pt(8), color=COLOR_SUB)


def _add_competitor_type_table(doc: Document, competitor_type: CompetitorType, group: list[Competitor]) -> None:
    label = f"{competitor_type.value} ({len(group)}개사)"
    _add_h2(doc, label)

    table = doc.add_table(rows=1, cols=len(_COMPETITOR_TABLE_HEADERS))
    table.autofit = False
    _apply_table_grid_fallback(table)

    header_row = table.rows[0]
    _mark_header_row(header_row)
    for cell, header_text, width in zip(header_row.cells, _COMPETITOR_TABLE_HEADERS, _COMPETITOR_TABLE_COL_WIDTHS_CM):
        cell.width = Cm(width)
        _shade_cell(cell, HEAD_BG_HEX)
        _set_cell_hv_borders(cell)
        _set_cell_margins(cell)
        _add_run(cell.paragraphs[0], header_text, bold=True, size=Pt(9.5), color=COLOR_WHITE)

    for idx, c in enumerate(group):
        row_cells = table.add_row().cells
        striped = idx % 2 == 1
        for cell, width in zip(row_cells, _COMPETITOR_TABLE_COL_WIDTHS_CM):
            cell.width = Cm(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            _set_cell_hv_borders(cell)
            _set_cell_margins(cell)
            if striped:
                _shade_cell(cell, LIGHT_HEX)

        name_p = row_cells[0].paragraphs[0]
        _add_run(name_p, c.name, bold=True, size=Pt(9.5), color=COLOR_PRIMARY)

        _add_run(row_cells[1].paragraphs[0], c.price or "(정보 없음)", size=Pt(8.5), color=COLOR_BODY)
        _feature_cell_text(row_cells[2], c.key_features)
        _target_remark_cell_text(row_cells[3], c)


def _add_competitor_footnote(doc: Document) -> None:
    _add_note(
        doc,
        "※ 대부분 경쟁사의 유통 채널·투자/매출 정보는 수집된 fact에 없어 표에서 제외했습니다"
        "(확보된 값만 '타깃·비고'에 표기). 가격 미공개 항목이 많아 가격 경쟁력 비교는 "
        "추가 조사가 필요합니다.",
    )


# ---------------------------------------------------------------------------
# 7. 05 종합 시사점 — 번호 카드
# ---------------------------------------------------------------------------
_SYNTHESIS_ITEM_RE = re.compile(r"(?:^|\n)\s*(\d+)[\.\)]\s*(.+?)(?=(?:\n\s*\d+[\.\)]\s*)|\Z)", re.S)


def _parse_synthesis_items(synthesis_text: str) -> list[tuple[str, str]]:
    """'1. 제목: 본문' 또는 '1. 본문' 형태의 번호 매긴 텍스트를 (제목, 본문) 튜플 목록으로
    분해한다. 번호 패턴이 전혀 없으면 빈 리스트를 반환해 호출부에서 평문 처리로 폴백한다."""
    items = []
    for match in _SYNTHESIS_ITEM_RE.finditer(synthesis_text.strip()):
        body_full = match.group(2).strip().replace("\n", " ")
        if ":" in body_full[:60]:
            title, _, rest = body_full.partition(":")
            items.append((title.strip(), rest.strip()))
        elif "：" in body_full[:60]:
            title, _, rest = body_full.partition("：")
            items.append((title.strip(), rest.strip()))
        else:
            words = body_full.split()
            title = " ".join(words[:6])
            items.append((title, body_full))
    return items


def _add_synthesis_card(doc: Document, number: int, title: str, body: str) -> None:
    col_widths_cm = [560 * _CM_PER_TWIP, 8746 * _CM_PER_TWIP]
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    _apply_table_grid_fallback(table)
    _set_table_grid_widths(table, col_widths_cm)
    badge_cell, content_cell = table.rows[0].cells
    badge_cell.width = Cm(col_widths_cm[0])
    content_cell.width = Cm(col_widths_cm[1])

    _shade_cell(badge_cell, ACCENT_HEX)
    _set_cell_hv_borders(badge_cell)
    _set_cell_margins(badge_cell)
    badge_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    bp = badge_cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.paragraph_format.space_before = Pt(0)
    bp.paragraph_format.space_after = Pt(0)
    _add_run(bp, str(number), bold=True, size=Pt(13), color=COLOR_WHITE)

    _shade_cell(content_cell, LIGHT_HEX)
    _set_cell_hv_borders(content_cell)
    _set_cell_margins(content_cell)
    content_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    title_p = content_cell.paragraphs[0]
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    _add_run(title_p, title, bold=True, size=Pt(10), color=COLOR_PRIMARY)
    body_p = content_cell.add_paragraph()
    body_p.paragraph_format.space_before = Pt(0)
    body_p.paragraph_format.space_after = Pt(0)
    _add_run(body_p, body, size=Pt(9), color=COLOR_BODY)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def _add_synthesis_section(doc: Document, synthesis_text: str) -> None:
    items = _parse_synthesis_items(synthesis_text)
    if not items:
        _add_body_paragraph(doc, synthesis_text)
        return
    for idx, (title, body) in enumerate(items, start=1):
        _add_synthesis_card(doc, idx, title, body)


# ---------------------------------------------------------------------------
# 8. 06 출처 부록 — 3열 표
# ---------------------------------------------------------------------------
def _add_source_appendix(doc: Document, fact_ids: list[str]) -> None:
    if not fact_ids:
        _add_body_paragraph(doc, "이 보고서가 직접 인용한 fact가 없습니다.")
        return

    note_p = doc.add_paragraph()
    _add_run(
        note_p,
        "Multi-Agent가 웹 검색으로 수집한 fact 목록입니다. [등급/판정]은 출처의 신뢰 "
        "등급(2차·3차)과 채택 여부(채택·애매·기각)를 의미합니다.",
        size=Pt(8.5),
        color=COLOR_SUB,
    )

    headers = ["등급/판정", "수집 fact", "출처"]
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    _apply_table_grid_fallback(table)
    header_row = table.rows[0]
    _mark_header_row(header_row)
    for cell, header_text, width in zip(header_row.cells, headers, _SOURCE_TABLE_COL_WIDTHS_CM):
        cell.width = Cm(width)
        _shade_cell(cell, HEAD_BG_HEX)
        _set_cell_hv_borders(cell)
        _set_cell_margins(cell)
        _add_run(cell.paragraphs[0], header_text, bold=True, size=Pt(9.5), color=COLOR_WHITE)

    row_idx = 0
    for fid in fact_ids:
        fact = get_fact(fid)
        if fact is None:
            continue
        status = fact.verification_status.value if fact.verification_status else "검증 미실시"
        row_cells = table.add_row().cells
        striped = row_idx % 2 == 1
        row_idx += 1
        for cell, width in zip(row_cells, _SOURCE_TABLE_COL_WIDTHS_CM):
            cell.width = Cm(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            _set_cell_hv_borders(cell)
            _set_cell_margins(cell)
            if striped:
                _shade_cell(cell, LIGHT_HEX)

        grade_p = row_cells[0].paragraphs[0]
        _add_run(grade_p, f"{fact.source_tier.value}/{status}", bold=True, size=Pt(7.5), color=COLOR_ACCENT)

        _add_run(row_cells[1].paragraphs[0], fact.text, size=Pt(7.5), color=COLOR_BODY)

        source_p = row_cells[2].paragraphs[0]
        _add_run(source_p, fact.source_url, size=Pt(6.5), color=COLOR_SUB)
        date_p = row_cells[2].add_paragraph()
        _add_run(date_p, f"(조회일: {fact.retrieved_date})", size=Pt(6.5), color=COLOR_SUB)


# ---------------------------------------------------------------------------
# 조립
# ---------------------------------------------------------------------------
def export_to_docx(
    topic: str,
    target_market: str,
    created_date: str,
    exec_summary: str,
    market_sizing: Optional[MarketSizing],
    pestel_summaries: list[dict],
    competitors: list[Competitor],
    positioning_narrative: str,
    synthesis: str,
    source_fact_ids: list[str],
    output_path: Path,
) -> Path:
    """구조화된 Writer 산출물을 A4 Word 문서로 저장하고 저장 경로를 반환한다."""
    doc = Document()
    _set_a4(doc)
    _add_footer(doc, f"{topic} — {target_market} 기획서 초안")

    _add_cover(doc, topic, target_market, created_date)

    # 01 개요
    _add_h1(doc, 1, "개요 (Executive Summary)")
    _add_kpi_cards(doc, market_sizing)
    doc.add_paragraph()
    _add_body_paragraph(doc, exec_summary, emphasize=True)

    # 02 시장 규모 분석
    _add_h1(doc, 2, "시장 규모 분석 (TAM · SAM · SOM)")
    if market_sizing is None:
        _add_body_paragraph(doc, "시장조사 에이전트 실행 이력이 없어 시장 규모 데이터가 없습니다.")
    else:
        _add_market_sizing_table(doc, market_sizing)
        doc.add_paragraph()
        if market_sizing.discrepancy_flag:
            _add_warn_box(doc, "Top-down/Bottom-up 추정치가 10배 이상 차이납니다 — 가정치 재검토 필요")
        if market_sizing.assumptions:
            _add_h2(doc, "가정 및 근거")
            _add_bullets(doc, market_sizing.assumptions)

    # 03 PESTEL 환경분석
    _add_h1(doc, 3, "PESTEL 환경분석")
    if not pestel_summaries:
        _add_body_paragraph(doc, "PESTEL 분석 대상 fact가 없습니다.")
    else:
        for s in pestel_summaries:
            _add_pestel_axis(doc, s["axis"], s["narrative"], s.get("excluded_low_materiality_fact_ids") or [])

    # 04 경쟁사 분석
    _add_h1(doc, 4, "경쟁사 분석")
    if not competitors:
        _add_body_paragraph(doc, "식별된 경쟁사가 없습니다.")
    else:
        by_type: dict = {t: [] for t in _COMPETITOR_TYPE_ORDER}
        for c in competitors:
            by_type.setdefault(c.type, []).append(c)
        for competitor_type in _COMPETITOR_TYPE_ORDER:
            group = by_type.get(competitor_type, [])
            if group:
                _add_competitor_type_table(doc, competitor_type, group)
        _add_competitor_footnote(doc)
        doc.add_paragraph()
        _add_h2(doc, "경쟁 구도 종합")
        _add_body_paragraph(doc, positioning_narrative, emphasize=True)

    # 05 종합 시사점
    _add_h1(doc, 5, "종합 시사점")
    _add_synthesis_section(doc, synthesis)

    # 06 출처 부록
    _add_h1(doc, 6, "출처 부록")
    _add_source_appendix(doc, source_fact_ids)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
